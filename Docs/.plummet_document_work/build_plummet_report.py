from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

ROOT = Path(r"D:\UnityProjects\Plummet")
WORK = ROOT / "Docs" / ".plummet_document_work"
REFERENCE = WORK / "Rock_VS_Mine_reference.docx"
OUTPUT = ROOT / "Docs" / "Plummet_Project_Report.docx"
ASSET_DIR = WORK / "generated_assets"
ASSET_DIR.mkdir(parents=True, exist_ok=True)

VIOLET = "3D3A87"
VIOLET_LIGHT = "E9E8F6"
INK = "1E2230"
MID = "5C6170"
PALE = "F5F5F8"
WHITE = "FFFFFF"


def font(name="times.ttf", size=28):
    path = Path(r"C:\Windows\Fonts") / name
    return ImageFont.truetype(str(path), size) if path.exists() else ImageFont.load_default()


F_B = font("timesbd.ttf", 31)
F_R = font("times.ttf", 25)
F_S = font("times.ttf", 20)


def rounded_box(draw, xy, title, lines=(), fill="#F5F5F8", outline="#3D3A87"):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=18, fill=fill, outline=outline, width=3)
    draw.text((x1 + 22, y1 + 18), title, fill="#3D3A87", font=F_B)
    y = y1 + 66
    for line in lines:
        draw.text((x1 + 22, y), line, fill="#1E2230", font=F_S)
        y += 30


def arrow(draw, start, end, color="#5C6170", width=4):
    draw.line([start, end], fill=color, width=width)
    x2, y2 = end
    x1, y1 = start
    dx, dy = x2 - x1, y2 - y1
    length = max((dx * dx + dy * dy) ** 0.5, 1)
    ux, uy = dx / length, dy / length
    px, py = -uy, ux
    a = (x2 - ux * 18 + px * 9, y2 - uy * 18 + py * 9)
    b = (x2 - ux * 18 - px * 9, y2 - uy * 18 - py * 9)
    draw.polygon([end, a, b], fill=color)


def make_architecture(path):
    img = Image.new("RGB", (1500, 1000), "white")
    d = ImageDraw.Draw(img)
    d.text((50, 32), "PLUMMET RUNTIME ARCHITECTURE", fill="#1E2230", font=font("timesbd.ttf", 38))
    layers = [
        (110, 125, 1390, 285, "PRESENTATION & INPUT", ["Main menu • level selection • HUD", "Input System • touch controls • camera and screen fade"], "#ECEBFA"),
        (110, 335, 1390, 495, "GAMEPLAY DOMAIN", ["Player finite-state machine • level actions • avatar control", "Doors • hazards • moving/falling ground • rule modifiers"], "#F3F0FB"),
        (110, 545, 1390, 705, "ORCHESTRATION", ["VContainer composition • signal bus • scene and level lifecycle", "Catalog selection • progress • retry/advance • audio routing"], "#F6F6F9"),
        (110, 755, 1390, 915, "PLATFORM & DATA", ["PlayerPrefs • Firebase Analytics • Smartlook", "Google Mobile Ads • Play in-app updates • Unity/Android"], "#F8F8FA"),
    ]
    for x1, y1, x2, y2, title, lines, fill in layers:
        rounded_box(d, (x1, y1, x2, y2), title, lines, fill=fill)
    for y in (285, 495, 705):
        arrow(d, (750, y + 5), (750, y + 45))
    img.save(path, dpi=(180, 180))


def make_fsm(path):
    img = Image.new("RGB", (1500, 840), "white")
    d = ImageDraw.Draw(img)
    d.text((50, 30), "PLAYER STATE MODEL", fill="#1E2230", font=font("timesbd.ttf", 38))
    nodes = {"Idle": (110, 180, 390, 330), "Run": (610, 180, 890, 330), "Jump": (1110, 180, 1390, 330), "Locked": (360, 570, 690, 720), "Dead": (860, 570, 1190, 720)}
    for title, xy in nodes.items():
        rounded_box(d, xy, title.upper(), (), fill="#ECEBFA" if title not in ("Dead", "Locked") else "#F6F6F9")
    arrow(d, (390, 255), (610, 255)); arrow(d, (610, 285), (390, 285))
    arrow(d, (890, 235), (1110, 235)); arrow(d, (1110, 295), (890, 295))
    arrow(d, (250, 330), (470, 570)); arrow(d, (750, 330), (560, 570)); arrow(d, (1250, 330), (650, 610))
    arrow(d, (300, 330), (940, 570)); arrow(d, (760, 330), (1010, 570)); arrow(d, (1250, 330), (1100, 570))
    d.text((462, 405), "door / scripted control", fill="#5C6170", font=F_S)
    d.text((910, 430), "hazard / fatal fall", fill="#5C6170", font=F_S)
    d.text((475, 760), "Exit completes → level lifecycle advances", fill="#3D3A87", font=F_R)
    img.save(path, dpi=(180, 180))


def make_level_flow(path):
    img = Image.new("RGB", (1500, 920), "white")
    d = ImageDraw.Draw(img)
    d.text((50, 30), "LEVEL LIFECYCLE AND EVENT FLOW", fill="#1E2230", font=font("timesbd.ttf", 38))
    boxes = [
        (80, 160, 380, 315, "SELECT", ["Catalog index", "Unlocked progress"]),
        (460, 160, 760, 315, "LOAD", ["Instantiate prefab", "Apply level rules"]),
        (840, 160, 1140, 315, "PLAY", ["Player / puppet", "Actions & hazards"]),
        (1180, 160, 1460, 315, "EXIT", ["Lock avatar", "Door animation"]),
        (230, 580, 560, 735, "RETRY", ["Death signal", "Reload same level"]),
        (650, 580, 980, 735, "ADVANCE", ["Persist unlock", "Load next level"]),
        (1070, 580, 1400, 735, "COMPLETE", ["Final catalog entry", "Game-complete event"]),
    ]
    for x1, y1, x2, y2, title, lines in boxes:
        rounded_box(d, (x1, y1, x2, y2), title, lines)
    arrow(d, (380, 235), (460, 235)); arrow(d, (760, 235), (840, 235)); arrow(d, (1140, 235), (1180, 235))
    arrow(d, (980, 580), (1240, 315)); arrow(d, (1320, 315), (1235, 580))
    arrow(d, (1190, 315), (820, 580)); arrow(d, (650, 655), (560, 655))
    arrow(d, (390, 580), (920, 315)); arrow(d, (230, 655), (100, 315))
    d.text((75, 820), "Cross-cutting signals: level started • player died • level completed • game completed", fill="#3D3A87", font=F_R)
    img.save(path, dpi=(180, 180))


def make_action_model(path):
    img = Image.new("RGB", (1500, 820), "white")
    d = ImageDraw.Draw(img)
    d.text((50, 30), "DATA-DRIVEN LEVEL ACTION MODEL", fill="#1E2230", font=font("timesbd.ttf", 38))
    rounded_box(d, (80, 165, 420, 365), "ACTIVATION", ["Trigger zone", "Auto-start", "One-shot guard"], fill="#ECEBFA")
    rounded_box(d, (580, 135, 920, 395), "SEQUENCE", ["Sequential", "or parallel", "Cancellation-aware"], fill="#F3F0FB")
    rounded_box(d, (1080, 95, 1430, 435), "ACTIONS", ["Move / return / wait", "Disable object", "Switch control", "Invert input", "Gravity & jump rules", "Camera shake / loop"], fill="#F6F6F9")
    rounded_box(d, (390, 565, 1110, 740), "RUNTIME EFFECT", ["Geometry, avatar control, and physics rules change", "without a new controller script."], fill="#F8F8FA")
    arrow(d, (420, 265), (580, 265)); arrow(d, (920, 265), (1080, 265)); arrow(d, (1250, 435), (890, 565)); arrow(d, (620, 565), (260, 365))
    img.save(path, dpi=(180, 180))


make_architecture(ASSET_DIR / "architecture.png")
make_fsm(ASSET_DIR / "player_fsm.png")
make_level_flow(ASSET_DIR / "level_flow.png")
make_action_model(ASSET_DIR / "action_model.png")

shutil.copyfile(REFERENCE, OUTPUT)
doc = Document(OUTPUT)
body = doc._element.body
for child in list(body):
    body.remove(child)
body.append(OxmlElement("w:sectPr"))


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=110, start=110, bottom=110, end=110):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v)); node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:tblHeader"); el.set(qn("w:val"), "true"); tr_pr.append(el)


def set_table_borders(tbl, color="B8BAC4", size="6"):
    tbl_pr = tbl._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}"); el.set(qn("w:val"), "single"); el.set(qn("w:sz"), size); el.set(qn("w:color"), color); borders.append(el)
    tbl_pr.append(borders)


def set_run_font(run, size=12, bold=None, italic=None, color=INK, name="Times New Roman"):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic
    if color: run.font.color.rgb = RGBColor.from_string(color)


def page_num_format(section, fmt="decimal", start=1):
    sect_pr = section._sectPr
    old = sect_pr.find(qn("w:pgNumType"))
    if old is not None: sect_pr.remove(old)
    node = OxmlElement("w:pgNumType")
    node.set(qn("w:fmt"), fmt); node.set(qn("w:start"), str(start)); sect_pr.append(node)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = " PAGE "
    separate = OxmlElement("w:fldChar"); separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t"); text.text = "1"
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run(); run._r.extend([begin, instr, separate, text, end]); set_run_font(run, 10, color=MID)


def setup_section(section, footer=True, fmt="decimal", start=1):
    section.page_width = Cm(21.0); section.page_height = Cm(29.7)
    section.left_margin = Inches(1.18); section.right_margin = Inches(0.98)
    section.top_margin = Inches(0.94); section.bottom_margin = Inches(0.86)
    section.header_distance = Inches(0.35); section.footer_distance = Inches(0.4)
    section.header.is_linked_to_previous = False; section.footer.is_linked_to_previous = False
    for container in (section.header._element, section.footer._element):
        for child in list(container):
            container.remove(child)
    page_num_format(section, fmt, start)
    if footer: add_page_field(section.footer.add_paragraph())


styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Times New Roman"; normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman"); normal.font.size = Pt(12)
normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; normal.paragraph_format.line_spacing = 1.45; normal.paragraph_format.space_after = Pt(7)
for name, size, before, after in (("Title", 20, 0, 10), ("Heading 1", 14, 12, 8), ("Heading 2", 12, 9, 5), ("Heading 3", 11, 7, 4)):
    style = styles[name] if name in [s.name for s in styles] else styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = "Times New Roman"; style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(size); style.font.bold = True; style.font.color.rgb = RGBColor.from_string(INK)
    style.paragraph_format.space_before = Pt(before); style.paragraph_format.space_after = Pt(after); style.paragraph_format.keep_with_next = True
if "Figure Caption" not in [s.name for s in styles]:
    cap_style = styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
else:
    cap_style = styles["Figure Caption"]
cap_style.font.name = "Times New Roman"; cap_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
cap_style.font.size = Pt(10); cap_style.font.italic = True
cap_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER; cap_style.paragraph_format.space_after = Pt(8)


def p(text="", *, align=None, bold=False, italic=False, size=12, style=None, before=0, after=7, keep=False):
    para = doc.add_paragraph(style=style)
    if align is not None: para.alignment = align
    para.paragraph_format.space_before = Pt(before); para.paragraph_format.space_after = Pt(after)
    para.paragraph_format.line_spacing = 1.45; para.paragraph_format.keep_with_next = keep
    if text:
        run = para.add_run(text); set_run_font(run, size, bold, italic)
    return para


def front_title(text):
    return p(text.upper(), align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14, after=18, keep=True)


def chapter(number, title, page_break=True):
    if page_break:
        doc.add_page_break()
    h = doc.add_paragraph(style="Heading 1"); h.alignment = WD_ALIGN_PARAGRAPH.CENTER; h.paragraph_format.space_after = Pt(15)
    run = h.add_run(f"CHAPTER {number}"); set_run_font(run, 14, True)
    run.add_break()
    title_run = h.add_run(title.upper()); set_run_font(title_run, 14, True)


def heading(text, level=2):
    h = doc.add_paragraph(style=f"Heading {level}"); run = h.add_run(text); set_run_font(run, 12 if level == 2 else 11, True); return h


def bullet(text, level=0):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.28 + level * 0.22); para.paragraph_format.first_line_indent = Inches(-0.18)
    para.paragraph_format.space_after = Pt(4); para.paragraph_format.line_spacing = 1.25
    run = para.add_run("• " + text); set_run_font(run, 11); return para


def number_item(text):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.33); para.paragraph_format.first_line_indent = Inches(-0.2); para.paragraph_format.space_after = Pt(5)
    run = para.add_run("– " + text); set_run_font(run, 11); return para


def table(headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers)); t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.autofit = False; set_table_borders(t)
    hdr = t.rows[0]; set_repeat_table_header(hdr)
    for idx, val in enumerate(headers):
        c = hdr.cells[idx]; set_cell_shading(c, VIOLET); c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER; set_cell_margins(c)
        if widths: c.width = Inches(widths[idx])
        pp = c.paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER; rr = pp.add_run(str(val)); set_run_font(rr, font_size, True, color=WHITE)
    for ridx, row in enumerate(rows):
        cells = t.add_row().cells
        for idx, val in enumerate(row):
            c = cells[idx]
            if ridx % 2: set_cell_shading(c, PALE)
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER; set_cell_margins(c)
            if widths: c.width = Inches(widths[idx])
            pp = c.paragraphs[0]; pp.paragraph_format.space_after = Pt(0); pp.paragraph_format.line_spacing = 1.1
            rr = pp.add_run(str(val)); set_run_font(rr, font_size)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return t


def figure(path, caption, width=6.3):
    para = doc.add_paragraph(); para.alignment = WD_ALIGN_PARAGRAPH.CENTER; para.paragraph_format.space_after = Pt(2); para.paragraph_format.keep_with_next = True
    shape = para.add_run().add_picture(str(path), width=Inches(width))
    shape._inline.docPr.set("descr", caption)
    shape._inline.docPr.set("title", caption.split(":", 1)[-1].strip())
    cp = doc.add_paragraph(style="Figure Caption"); run = cp.add_run(caption); set_run_font(run, 10, italic=True)


def signature_table(items):
    t = doc.add_table(rows=2, cols=len(items)); t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.autofit = False
    for i, (name, role) in enumerate(items):
        c = t.cell(0, i); c.width = Inches(2.1); pp = c.paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(pp.add_run("________________________"), 10)
        c2 = t.cell(1, i); pp2 = c2.paragraphs[0]; pp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(pp2.add_run(name + "\n"), 11, True); set_run_font(pp2.add_run(role), 10)
    for row in t.rows:
        for c in row.cells:
            tc_pr = c._tc.get_or_add_tcPr(); borders = OxmlElement("w:tcBorders")
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                el = OxmlElement(f"w:{edge}"); el.set(qn("w:val"), "nil"); borders.append(el)
            tc_pr.append(borders)
    return t


def add_toc_field():
    para = doc.add_paragraph(); run = para.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    separate = OxmlElement("w:fldChar"); separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t"); text.text = "Update field to generate table of contents."
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


# Cover
setup_section(doc.sections[0], footer=False)
doc.sections[0].top_margin = Inches(0.72)
p("Project for 6th Semester of Bachelor of Information Technology", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=13, after=16)
p("PLUMMET", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=24, after=4)
p("A 2D Puzzle-Platformer for Android", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=15, after=10)
logo_p = doc.add_paragraph(); logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
logo_shape = logo_p.add_run().add_picture(str(WORK / "reference_media" / "image1.png"), width=Inches(1.58))
logo_shape._inline.docPr.set("descr", "Purbanchal University emblem")
logo_shape._inline.docPr.set("title", "Purbanchal University")
p("Submitted by", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12, after=5)
table(["Student", "Class Roll No."], [("Mausham Neupane", "360379"), ("Mohammad Aarman", "360381"), ("Tapendra Shahi", "360400")], widths=[3.4, 1.7], font_size=10)
p("KIST College of Information Technology", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=13, after=3)
p("Kamalpokhari, Kathmandu", align=WD_ALIGN_PARAGRAPH.CENTER, size=11, after=8)
p("Submitted to", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=11, after=2)
p("Faculty of Science and Technology\nPurbanchal University", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12, after=8)
p("August 2026", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12, after=0)

# Preliminary pages
prelim = doc.add_section(WD_SECTION.NEW_PAGE)
setup_section(prelim, footer=True, fmt="lowerRoman", start=1)
front_title("Student Declaration")
p("We hereby declare that the project report entitled “Plummet: A 2D Puzzle-Platformer for Android” is an original account of the software system developed and examined by us for the sixth semester of the Bachelor of Information Technology program. The report has been prepared from the current Unity project, its complete authored source code, scene and prefab configuration, project settings, and direct Play Mode observation. Material from the earlier Rock VS Mine report has been used only as an academic layout and structural reference.")
p("The work has not been submitted elsewhere for an academic award. All external frameworks, packages, services, and documentation consulted during development or preparation of this report are acknowledged in the references.")
table(["Student", "Registration No.", "Signature"], [("Mausham Neupane", "058-3-2-05288-2022", "________________"), ("Mohammad Aarman", "058-3-2-05289-2022", "________________"), ("Tapendra Shahi", "058-3-2-05310-2022", "________________")], widths=[2.35, 2.35, 1.6], font_size=10)
p("Date: ____________________", align=WD_ALIGN_PARAGRAPH.RIGHT, size=11, before=20)

doc.add_page_break(); front_title("Recommendation")
p("This is to certify that the project report entitled “Plummet: A 2D Puzzle-Platformer for Android” has been prepared by Mausham Neupane, Mohammad Aarman, and Tapendra Shahi under the academic requirements of the Bachelor of Information Technology program. The report presents the design, implementation, and verification of the submitted Unity project and is recommended for evaluation.")
p("The technical account is organized around observable project evidence: the runtime architecture, gameplay state model, data-driven level-action system, platform integrations, content catalog, and current validation status.")
p("", before=55, after=0)
signature_table([("Deepak Khadka", "Coordinator, B.I.T."), ("Bishal Khatri", "Project Supervisor")])

doc.add_page_break(); front_title("Certificate")
p("This is to certify that the project entitled “Plummet: A 2D Puzzle-Platformer for Android” has been examined and accepted as a partial fulfillment of the requirements for the Bachelor of Information Technology degree of Purbanchal University.")
p("The undersigned certify that the report describes the submitted project in a clear and professionally structured manner, subject to the regulations and final evaluation procedures of the University.")
p("", before=52, after=0)
signature_table([("Pankaj Jaiswal", "External Examiner"), ("Bishal Khatri", "Project Supervisor"), ("Deepak Khadka", "Coordinator, B.I.T.")])

doc.add_page_break(); front_title("Acknowledgement")
p("We express our sincere gratitude to KIST College of Information Technology and the Faculty of Science and Technology, Purbanchal University, for providing the academic setting in which this project was completed. We are especially thankful to our project supervisor, Mr. Bishal Khatri, and our program coordinator, Mr. Deepak Khadka, for their guidance, review, and encouragement.")
p("We also acknowledge the developers and maintainers of Unity, VContainer, UniTask, PrimeTween, Firebase, Google Mobile Ads, Smartlook, and the other packages used by the project. Their tools enabled us to concentrate on gameplay design, reliable scene orchestration, responsive input, and mobile deployment.")
p("Finally, we thank our classmates, friends, and family for testing, feedback, and continued support throughout the design and implementation of Plummet.")

doc.add_page_break(); front_title("Abstract")
p("Plummet is a two-dimensional puzzle-platformer developed in Unity for Android. The player navigates compact pixel-art levels by moving and jumping through hazards, falling platforms, moving geometry, rule-changing triggers, and exit doors. Its central design idea is that a level may alter the assumptions under which the player acts: controls can invert, gravity and jump behavior can change, and control can transfer to a puppet avatar. These changes are implemented through a reusable, serializable action framework rather than a separate controller for every level.")
p("The software uses a layered runtime architecture. VContainer establishes dependency injection scopes; a signal bus coordinates menu, player, and level events; additive scene loading separates bootstrap, menu, and gameplay responsibilities; and a player finite-state machine isolates idle, running, jumping, locked, and dead behavior. PlayerPrefs stores progress and audio settings. Firebase Analytics, Smartlook, Google Mobile Ads, and Google Play in-app updates provide platform capabilities behind service interfaces.")
p("This report is grounded in a complete inspection of the authored scripts, scenes, prefabs, package manifest, project settings, level catalog, and a Play Mode pass. The project contains 115 authored C# files totaling 7,956 lines, three build scenes, 57 prefabs, and 21 cataloged levels. A Level 22 prefab is present but not included in the runtime catalog. The system provides a maintainable mobile puzzle-platformer foundation with clear opportunities for automated testing, catalog validation, and production service configuration.", size=11.5)

doc.add_page_break(); front_title("Table of Contents")
add_toc_field()

doc.add_page_break(); front_title("List of Figures")
table(["No.", "Figure"], [("3.1", "Plummet runtime architecture"), ("3.2", "Player finite-state model"), ("3.3", "Level lifecycle and event flow"), ("3.4", "Data-driven level action model"), ("4.1", "Level 1 Play Mode capture"), ("4.2", "Representative hazard and platform composition")], widths=[0.75, 5.55], font_size=10)

doc.add_page_break(); front_title("Abbreviations")
table(["Abbreviation", "Meaning"], [("API", "Application Programming Interface"), ("DI", "Dependency Injection"), ("FSM", "Finite-State Machine"), ("FPS", "Frames Per Second"), ("SDK", "Software Development Kit"), ("SFX", "Sound Effects"), ("UI", "User Interface"), ("URP", "Universal Render Pipeline"), ("UX", "User Experience"), ("UPM", "Unity Package Manager")], widths=[1.45, 4.85], font_size=10)

# Main body
body_sec = doc.add_section(WD_SECTION.NEW_PAGE)
setup_section(body_sec, footer=True, fmt="decimal", start=1)
body_sec.bottom_margin = Inches(1.0)

chapter(1, "Introduction", page_break=False)
heading("1.1 Background")
p("Puzzle-platformers combine physical navigation with systems that require the player to infer rules, predict consequences, and execute precise actions. On a mobile device, this creates a demanding interaction problem: the game must remain visually legible, respond quickly to touch input, and communicate failure without interrupting the rapid retry rhythm that makes short challenge levels engaging.")
p("Plummet addresses this problem through a compact 2D pixel-art presentation and a modular gameplay architecture. The player moves between an entrance and an exit while avoiding spikes, saws, fatal falls, and unstable geometry. Some levels deliberately transform the interaction model by inverting controls, changing gravity or jump allowances, transferring control to a puppet, or executing a timed sequence of environmental actions.")
heading("1.2 Problem Statement")
p("A conventional platformer can hard-code each challenge directly into a level script, but that approach becomes difficult to extend and verify as the number of levels grows. Plummet therefore requires a system that can express unusual level behavior from reusable parts, preserve consistent player rules, support fast death-and-retry loops, and integrate mobile platform services without coupling them to gameplay code.")
heading("1.3 Objectives")
number_item("Design a responsive movement and jump model suitable for desktop testing and touch-screen play.")
number_item("Represent player behavior with explicit states so movement, death, locking, and scripted control remain predictable.")
number_item("Create a reusable level-action framework for movement, waits, object disabling, physics-rule changes, control switching, camera shake, and loops.")
number_item("Build a persistent level-selection and progression system for a catalog of short puzzle stages.")
number_item("Separate gameplay from advertising, analytics, recording, audio, saving, and update services through interfaces.")
number_item("Validate the submitted project through structural inspection and direct Play Mode observation.")
heading("1.4 Scope")
p("The project scope includes the Android game client, bootstrap/menu/gameplay scenes, player and puppet control, level catalog and instantiation, environmental actions, audio settings, local progress, analytics events, mobile advertisements, session recording integration, and in-app update support. Server-hosted gameplay, online accounts, multiplayer, cloud saves, and a remote level editor are outside the present scope.")

chapter(2, "Requirements Analysis and Feasibility")
heading("2.1 Functional Requirements")
table(["ID", "Requirement", "Verified implementation evidence"], [
    ("FR-01", "Start and navigate the game", "Bootstrap loads the configured start scene; menu exposes Play, Music, and SFX controls."),
    ("FR-02", "Select unlocked levels", "LevelSelection uses catalog entries and saved highest-unlocked progress."),
    ("FR-03", "Move and jump", "Input System actions feed the player state machine and on-screen controls."),
    ("FR-04", "Die and retry quickly", "Hazards publish death; the runtime reloads the same level after the death sequence."),
    ("FR-05", "Complete and advance", "Exit door locks the avatar, animates closure, persists progress, and requests the next level."),
    ("FR-06", "Alter level rules", "Serializable actions can invert input, adjust gravity/jump rules, and switch the controlled avatar."),
    ("FR-07", "Persist preferences", "PlayerPrefs stores music, sound-effect, and level-progress values."),
    ("FR-08", "Provide mobile services", "Ads, analytics, Smartlook, and in-app updates are wrapped behind runtime services."),
], widths=[0.65, 2.2, 3.45], font_size=8.7)
heading("2.2 Non-Functional Requirements")
table(["Quality", "Requirement and design response"], [
    ("Responsiveness", "Input is read through a single input contract; movement and physics adjustments are applied in dedicated player states."),
    ("Maintainability", "Interfaces, dependency injection, state classes, signal-driven orchestration, and catalog assets limit direct scene dependencies."),
    ("Extensibility", "New level behavior can be assembled from serializable actions and sequences without rewriting the central level runtime."),
    ("Visual clarity", "Pixel-art scenes use a 2D rendering stack, constrained framing, and a border overlay for aspect-ratio differences."),
    ("Resilience", "Analytics queues events until initialization; service implementations isolate platform-specific dependencies."),
    ("Usability", "Large on-screen directional and jump controls support the primary mobile interaction path."),
], widths=[1.45, 4.85], font_size=9.3)
heading("2.3 Feasibility")
p("Technical feasibility is supported by Unity 6000.3.18f1, the Universal Render Pipeline, the Input System, and established packages for asynchronous work, dependency injection, tweening, and platform services. The project builds around a small number of scenes and instantiates level prefabs from a catalog, keeping content iteration separate from scene topology.")
p("Operational feasibility is strengthened by short self-contained levels, immediate failure feedback, saved unlock progress, and familiar touch controls. Economic feasibility is supported by the use of broadly available Unity tooling and optional advertising, while the present implementation avoids a required custom backend. Schedule feasibility is improved by the reusable level-action framework and editor utilities such as the sprite-combining tool.")
heading("2.4 Constraints and Assumptions")
bullet("The shipping target is Android; some external services require a physical device, valid credentials, or Play Store context for complete verification.")
bullet("The current catalog exposes 21 levels even though a Level 22 prefab exists in the project.")
bullet("The inspected input asset defines keyboard actions (A, D, and Space) in addition to scene-based touch controls.")
bullet("The current report verifies behavior and configuration but does not claim store-release certification or device-matrix coverage.")

chapter(3, "System Methodology and Design")
heading("3.1 Development Method")
p("Plummet follows an iterative, component-oriented game-development method. Core movement and scene flow form the stable foundation; each puzzle stage composes prefabs, triggers, actions, and level rules on top of that foundation. Reusable contracts are defined first, platform implementations are injected at the composition root, and gameplay coordination occurs through typed signals. This keeps experimentation local while preserving deterministic lifecycle boundaries.")
heading("3.2 Runtime Architecture")
p("The software can be understood as four cooperating layers. Presentation and input collect player intent and display menus, controls, camera effects, and fades. The gameplay domain owns player states, avatar behavior, hazards, doors, and level actions. Orchestration selects and loads content, transfers control, saves progress, and routes signals. Platform and data adapters communicate with PlayerPrefs and external mobile SDKs.")
figure(ASSET_DIR / "architecture.png", "Figure 3.1: Plummet runtime architecture", width=5.95)
heading("3.3 Composition and Scene Lifecycle")
p("RootLifetimeScope is the application composition root. It registers the level catalog, input, audio library, settings, signal bus, scene loader, level selection, persistence, advertisements, analytics, Smartlook, in-app updates, and bootstrap entry point. Scene-specific lifetime scopes extend that root for menu and gameplay dependencies. Bootstrap, MainMenu, and Gameplay are the three enabled build scenes; additive loading and a screen fader hide transitions and allow the root services to remain available.")
heading("3.4 Player Finite-State Machine")
p("The player controller delegates behavior to five explicit states. Idle waits for movement or jump input, Run applies horizontal movement, Jump manages airborne motion and enhanced falling, Locked prevents ordinary control during scripted events, and Dead prevents further play until the level runtime retries. Default inspected values are movement speed 6, jump force 12, and fall multiplier 1; level rules may modify the effective jump and gravity behavior.")
figure(ASSET_DIR / "player_fsm.png", "Figure 3.2: Player finite-state model", width=6.45)
heading("3.5 Level Lifecycle")
p("LevelSelection identifies the active catalog entry. The runtime instantiates that level prefab through VContainer, applies its rules, spawns or binds the avatar, and records the level-start event. A death signal increments the retry path and reloads the same selection; an exit sequence records completion, unlocks progress, and advances. Reaching the final catalog entry publishes a game-complete event. Interstitial display is attempted after a randomized interval of seven to ten deaths when an advertisement is ready.")
figure(ASSET_DIR / "level_flow.png", "Figure 3.3: Level lifecycle and event flow", width=6.45)
heading("3.6 Data-Driven Level Actions")
p("The central extensibility mechanism is a serializable action hierarchy. Actions can run sequentially or in parallel and are commonly activated by a Player-tag trigger zone or by automatic start. Implemented operations include moving an object, returning it, waiting, disabling an object, camera shaking, setting or flipping inverted controls, changing gravity and jump parameters, enabling or disabling jumping, switching the controlled avatar, applying or restoring level rules, and looping a sequence.")
figure(ASSET_DIR / "action_model.png", "Figure 3.4: Data-driven level action model", width=6.45)
heading("3.7 Data and Integration Design")
table(["Concern", "Primary mechanism", "Boundary"], [
    ("Progress", "Highest unlocked level in PlayerPrefs", "ILevelProgress / ILevelSelection"),
    ("Audio", "Music and SFX settings plus AudioLibrary", "IAudioService"),
    ("Analytics", "Queued Firebase events", "IAnalyticsService"),
    ("Session insight", "Mobile Smartlook integration", "ISmartlookService"),
    ("Advertising", "Banner, interstitial, rewarded units", "IAdsService"),
    ("Updates", "Google Play in-app update flow", "InAppUpdate service"),
    ("Coordination", "Typed publish/subscribe messages", "ISignalBus"),
], widths=[1.35, 2.55, 2.4], font_size=9.2)

chapter(4, "Implementation")
heading("4.1 Technology Stack")
table(["Technology", "Inspected version/source", "Role"], [
    ("Unity", "6000.3.18f1", "Editor, engine, scenes, serialization, Android build target"),
    ("URP", "17.3.0", "2D rendering pipeline"),
    ("Input System", "1.19.0", "Keyboard and touch-control input actions"),
    ("VContainer", "Git dependency, tag 1.18.0", "Dependency injection and lifetime scopes"),
    ("UniTask", "UPM/Git package", "Allocation-conscious asynchronous workflows"),
    ("PrimeTween", "Local package archive", "Screen, camera, and door/player tweening"),
    ("Firebase Analytics", "Project plugin", "Gameplay event reporting"),
    ("Google Mobile Ads", "Project plugin", "Banner, interstitial, and rewarded advertisements"),
], widths=[1.45, 2.05, 2.8], font_size=8.9)
heading("4.2 Source Organization")
p("The inspected authored code contains 115 C# files and 7,956 lines under Assets/DayZeroGames/Scripts. Of these, 106 are runtime scripts, eight are editor scripts, and one is the generated input wrapper. The code is organized by contracts, installers/lifetime scopes, player and puppet controllers, level runtime/actions, menu UI, services, audio, camera effects, and editor tooling.")
table(["Module", "Representative responsibilities"], [
    ("Contracts", "Ads, analytics, audio, input, saving, scenes, state, avatars, progress, and signals"),
    ("Bootstrap & DI", "Global registrations, startup entry point, update service, scene scopes"),
    ("Player", "FSM, movement, jump, death, collision and hazard responses"),
    ("Level Runtime", "Catalog selection, prefab instantiation, spawn, retry, next-level and completion"),
    ("Level Actions", "Sequence execution, movement, waits, rule changes, control switching and loops"),
    ("Presentation", "Main menu, level buttons, on-screen controls, screen fade and camera border"),
    ("Services", "PlayerPrefs, Firebase, Smartlook, ads, audio, and scene loading"),
    ("Editor", "Sprite combiner and supporting inspectors/utilities"),
], widths=[1.45, 4.85], font_size=9.2)
heading("4.3 Player and Avatar Control")
p("Input is exposed through an abstraction rather than queried throughout gameplay. The normal player and puppet controller consume the same rule source, allowing a level to transfer control without changing the overall lifecycle. Horizontal movement is physics-based, jumping respects maximum air jumps and jump-force multipliers, and falling can be accelerated. When an exit is reached, the avatar is locked, tweened into the door, faded, and followed by a close-door sequence before advancement.")
heading("4.4 Gameplay Evidence")
p("The following capture was taken from the submitted project in Play Mode. It shows the pixel-art composition, entrance and exit doors, player avatar, layered platforms, spike hazards, and large mobile controls. The framing also demonstrates the intended high-contrast separation between safe surfaces, hazards, and the dark background.")
figure(WORK / "game_captures" / "gameplay-level-1.png", "Figure 4.1: Level 1 Play Mode capture", width=6.45)
p("Levels combine direct objects with nested prefabs. Across the level set, the authored action sequences use moving and returning platforms, timed waits, object removal, changes to gravity and control direction, puppet control, parallel execution, and repeated loops. Later prefabs increase the density of moving or falling geometry and hazards while reusing the same runtime contracts.")
figure(WORK / "game_captures" / "gameplay-level-2.png", "Figure 4.2: Representative hazard and platform composition", width=6.45)
heading("4.5 Platform Services")
p("Platform dependencies are not called directly from the player or level code. The analytics service queues gameplay events until Firebase becomes ready and emits player_died, level_started, level_completed, and game_completed data. The advertising service manages banner, interstitial, and rewarded formats. Smartlook is restricted to supported mobile builds, and the Play in-app update integration handles Android distribution updates. This separation allows gameplay to remain testable in the Editor when external services are unavailable.")

chapter(5, "Verification and Results")
heading("5.1 Verification Method")
p("Verification combined static and dynamic evidence. Static inspection covered every authored C# script, build settings, tags and layers, the package manifest, ProjectSettings, LevelCatalog.asset, scene hierarchy, and all 22 level prefabs. Dynamic inspection entered Play Mode, observed the MainMenu and Gameplay scenes, loaded representative level prefabs through the runtime, captured the visible game view, and examined the Unity Console.")
heading("5.2 Structural Results")
table(["Item", "Observed result", "Status"], [
    ("Authored C# source", "115 files / 7,956 lines inspected", "Verified"),
    ("Build scenes", "Bootstrap, MainMenu, Gameplay", "Verified"),
    ("Level catalog", "21 ordered entries, Level 1 through Level 21", "Verified"),
    ("Additional content", "Level 22 prefab exists outside the catalog", "Attention"),
    ("Prefabs", "57 total; 22 top-level level prefabs", "Verified"),
    ("Rendering/input", "URP 17.3.0 and Input System 1.19.0", "Verified"),
    ("Android metadata", "Product Plummet, company DayZero Games, bundle version 1.6/code 7", "Verified"),
], widths=[1.35, 3.85, 1.1], font_size=9)
heading("5.3 Play Mode Results")
table(["Check", "Evidence", "Outcome"], [
    ("Startup", "Bootstrap transitioned to the configured menu scene", "Pass"),
    ("Gameplay load", "Representative level prefabs instantiated and rendered", "Pass"),
    ("Player presentation", "Player, doors, platforms, hazards, and mobile controls visible", "Pass"),
    ("Console", "No script exception or compilation error observed in the verification pass", "Pass"),
    ("External configuration", "Editor automation/audio-listener timing and Firebase database-URL warnings appeared", "Review"),
], widths=[1.25, 4.15, 0.9], font_size=9.1)
p("The observed warnings did not prevent representative levels from loading.")
heading("5.4 Level Content Coverage")
table(["Levels", "Observed composition emphasis"], [
    ("1", "Control inversion, reduced gravity, puppet switching, moving/returning objects, spikes and falling ground"),
    ("2–3", "Static hazard introduction followed by triggered movement, waits, and object disabling"),
    ("4–6", "Dense parallel fall/move sequences, fake or falling geometry, and moving platforms"),
    ("7–10", "Nested/static challenge patterns with spikes, repeated movement, returns, waits, and timed disabling"),
    ("11–16", "Larger trigger networks, parallel sequences, moving platforms, spikes, and falling-ground combinations"),
    ("17–19", "Primarily nested prefab composition with falling-ground and spike challenges"),
    ("20–21", "Multiple object disables and movement sequences; Level 21 uses several parallel groups"),
    ("22", "Prefab includes a saw hazard but is not referenced by the runtime catalog"),
], widths=[1.0, 5.3], font_size=9.1)
heading("5.5 Limitations and Risks")
bullet("No authored automated test suite was identified; current assurance depends primarily on structural review and Play Mode testing.")
bullet("LevelCatalog and level-prefab inventory can diverge, as demonstrated by Level 22; an editor validation rule would prevent accidental omission.")
bullet("Advertising, analytics, recording, and update behavior require device builds and valid production credentials for end-to-end verification.")
bullet("The current local save uses PlayerPrefs and does not provide cloud synchronization, account portability, or tamper resistance.")
bullet("A broader device and aspect-ratio matrix is required before release-quality claims about performance and touch ergonomics.")
doc.add_page_break()
heading("5.6 Recommended Next Steps")
number_item("Add EditMode tests for catalog integrity, progress bounds, and rule composition, plus PlayMode tests for death/retry and exit/advance flows.")
number_item("Decide whether Level 22 should be added to the catalog or intentionally archived, and enforce the decision with editor validation.")
number_item("Run Android device tests for advertisements, Firebase events, Smartlook capture, in-app updates, audio focus, and lifecycle pause/resume.")
number_item("Establish frame-time, memory, loading-time, crash-free-session, and level-completion targets for release monitoring.")

chapter(6, "Conclusion")
p("Plummet is a modular foundation for a mobile 2D puzzle-platformer. It provides a reusable runtime for loading content, coordinating typed events, controlling multiple avatars, changing physics and input rules, saving progress, and integrating mobile services. The player state machine and data-driven action system are its strongest architectural decisions because they isolate behavior and let designers assemble unusual challenges without duplicating the central controller.")
p("Inspection verified 115 authored C# files, three build scenes, 57 prefabs, and 21 cataloged levels. Representative Play Mode runs rendered without a script exception. The main readiness gaps are clear: resolve the uncataloged Level 22 prefab, introduce automated tests, and validate mobile service integrations on a production-like Android device and distribution track.")
p("With those additions, Plummet can progress from a well-structured academic game project to a more reliably testable and release-ready product. Its separation of gameplay, orchestration, data, and platform concerns provides a sound base for that next phase.")

chapter(7, "References")
refs = [
    "Unity Technologies. Unity Manual: Input System. https://docs.unity3d.com/Manual/com.unity.inputsystem.html",
    "Unity Technologies. Universal Render Pipeline 2D documentation. https://docs.unity3d.com/6000.0/Manual/urp/2d-index.html",
    "HadashiA. VContainer: dependency injection library for Unity. https://github.com/hadashiA/VContainer",
    "Cysharp. UniTask: efficient async/await integration for Unity. https://github.com/Cysharp/UniTask",
    "Kyrylo Kuzyk. PrimeTween for Unity. https://github.com/KyryloKuzyk/PrimeTween",
    "Google. Get started with Firebase for Unity. https://firebase.google.com/docs/unity/setup",
    "Google. Google Mobile Ads Unity plug-in quick start. https://developers.google.com/admob/unity/quick-start",
    "Android Developers. Support in-app updates in Unity. https://developer.android.com/guide/playcore/in-app-updates/unity",
    "DayZero Games. Plummet Unity project source, assets, scenes, prefabs, package manifest, and project settings. Inspected 24 August 2026.",
]
for i, ref in enumerate(refs, 1):
    para = doc.add_paragraph(); para.alignment = WD_ALIGN_PARAGRAPH.LEFT; para.paragraph_format.left_indent = Inches(0.28); para.paragraph_format.first_line_indent = Inches(-0.28); para.paragraph_format.space_after = Pt(7); para.paragraph_format.line_spacing = 1.2
    set_run_font(para.add_run(f"[{i}] {ref}"), 11)

chapter("A", "Appendix: Codebase Traceability")
heading("A.1 Runtime Inventory")
table(["Area", "Evidence reviewed", "Purpose in the implemented system"], [
    ("Contracts", "All interface and signal definitions", "Stable boundaries for services, state, avatars, saving, scenes, and events"),
    ("Composition roots", "Root, menu, gameplay and level scopes", "Object lifetime, dependency registration, and entry-point startup"),
    ("Player/Puppet", "Controllers, states, rules, collisions", "Movement, jumping, death, scripted locking, and alternate-avatar play"),
    ("Level core", "Catalog, runtime, selection, progress, exits", "Content selection, instantiation, retry, advancement, and completion"),
    ("Action system", "All concrete actions and sequence runners", "Reusable changes to geometry, time, input, physics, and control"),
    ("Menu/UI", "Menu and dynamic level-button scripts", "Start flow, settings, and progress-aware level navigation"),
    ("Platform services", "Ads, analytics, recording, update, save, audio", "External SDK isolation and persisted local state"),
    ("Editor tooling", "Sprite-combining editor utilities", "Rasterized composite sprites, order preservation, trimming, and collider generation"),
], widths=[1.35, 2.15, 2.8], font_size=8.2)
heading("A.2 Project Inventory")
table(["Artifact class", "Count / value", "Inspection note"], [
    ("Authored C# files", "115", "106 runtime, 8 editor, 1 generated input wrapper"),
    ("Authored C# lines", "7,956", "Counted under Assets/DayZeroGames/Scripts"),
    ("Scenes", "3 enabled", "Bootstrap, MainMenu, Gameplay"),
    ("Prefabs", "57", "Includes 22 top-level level prefabs"),
    ("Cataloged levels", "21", "Lvl_1 through Lvl_21 in LevelCatalog.asset"),
    ("PNG assets", "59", "Pixel-art and interface imagery"),
    ("Audio assets", "4", "BackgroundMusic, Death, ExitDoorReached, UIButtonPressed"),
    ("Tags", "9 built-in/custom", "Includes Player, Obstacles, FallingGround, Finish"),
], widths=[1.55, 1.35, 3.4], font_size=8.6)

chapter("B", "Appendix: Level-Action Coverage")
table(["Level", "Direct serialized action evidence"], [
    ("1", "Disable ×1; Loop ×1; Move ×2; Return ×1; SetGravity ×1; SetInvert ×1; SwitchControl ×1"),
    ("2", "No direct actions; static spike challenge"),
    ("3", "Disable ×1; Move ×2; Return ×1; Wait ×1"),
    ("4", "Disable ×21; Move ×21; Wait ×10; eleven parallel sequences"),
    ("5", "Disable ×1; Move ×1"),
    ("6", "Disable ×2; Move ×6"),
    ("7", "No direct action payload identified in the top-level prefab"),
    ("8", "Disable ×2; Move ×6; Return ×3; Wait ×5"),
    ("9", "Disable ×1; Move ×8; Return ×2; Wait ×5"),
    ("10", "Disable ×1; Move ×4"),
    ("11", "Disable ×7; Move ×11; Wait ×9"),
    ("12", "Disable ×1; Move ×4; Wait ×1"),
    ("13", "No direct actions; falling-ground composition"),
    ("14", "Disable ×1; Move ×6; Return ×1; Wait ×2"),
    ("15", "Move ×2"),
    ("16", "Disable ×7; Move ×7; five parallel sequences"),
    ("17–19", "No direct top-level actions; nested prefabs with falling-ground/spike patterns"),
    ("20", "Disable ×3; Move ×3"),
    ("21", "Disable ×3; Move ×9; Wait ×2; four sequences, three parallel"),
    ("22", "No direct actions; saw hazard; prefab is not cataloged"),
], widths=[0.9, 5.4], font_size=8.8)
p("Note: “direct” refers to serialized actions located in the inspected top-level level prefab. Nested prefab behavior may add further runtime composition.", italic=True, size=10, after=0)

settings = doc.settings._element
update = settings.find(qn("w:updateFields"))
if update is None:
    update = OxmlElement("w:updateFields"); settings.append(update)
update.set(qn("w:val"), "true")
doc.save(OUTPUT)
print(OUTPUT)
